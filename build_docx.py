"""
Resume — clean single-column, accent-line template.
Font: Times New Roman, black throughout. Target: 1 page.
"""

import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK  = RGBColor(0, 0, 0)
FONT   = "Times New Roman"
BODY   = 11
NAME_S = 16
TITLE_S = 11
CONT_S = 10
HEAD_S = 11
USABLE_W_TWIPS = int(Inches(7.5).pt * 20)

with open("/Users/susanokevinamalamahilmaran/Downloads/Projects/Resume/resume.json") as f:
    data = json.load(f)

MONTHS = ["Jan","Feb","Mar","Apr","May","Jun",
          "Jul","Aug","Sep","Oct","Nov","Dec"]

def fmt_date(d):
    if not d: return ""
    y, m = d.split("-")
    return f"{MONTHS[int(m)-1]}. {y}"


# ── helpers ──────────────────────────────────────────────────────

def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin   = Inches(0.5)
    sec.right_margin  = Inches(0.5)
    s = doc.styles["Normal"]
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after  = Pt(0)
    s.font.name      = FONT
    s.font.color.rgb = BLACK
    s.font.size      = Pt(BODY)
    return doc


def run(p, text, bold=False, size=BODY, italic=False):
    r = p.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.name   = FONT
    r.font.size   = Pt(size)
    r.font.color.rgb = BLACK
    return r


def para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    return p


def add_border(p, side="bottom", sz="8", color="000000", val="single", space="2"):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el   = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   val)
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


def section_head(doc, title, sb=10):
    p = para(doc, sb=sb, sa=1)
    add_border(p, sz="4", space="1")
    run(p, title, bold=True, size=HEAD_S)


def right_tab_para(doc, left, right, lb=True, rb=False, sb=0, sa=0,
                   lsize=BODY, rsize=BODY):
    p = para(doc, sb=sb, sa=sa)
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(USABLE_W_TWIPS))
    tabs.append(tab)
    pPr.append(tabs)
    run(p, left,          bold=lb, size=lsize)
    run(p, "\t" + right,  bold=rb, size=rsize)


def bullet(doc, text):
    p = para(doc, sa=2)
    p.paragraph_format.left_indent       = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run(p, "•  " + text, size=BODY)


# ── Build ────────────────────────────────────────────────────────

doc = setup_doc()

# ── HEADER ───────────────────────────────────────────────────────
p = para(doc, sa=2)
run(p, data["personal"]["name"], bold=True, size=NAME_S)

p = para(doc, sa=2)
run(p, "Data Analyst  &  AI Engineer", italic=True, size=TITLE_S)

# Contact line with thick bottom rule
p = para(doc, sa=0)
add_border(p, sz="12", space="3")
contact = (f"{data['personal']['email']}  |  "
           f"linkedin.com/in/susano-kevin  |  github.com/SusanoKevin  |  Milwaukee, WI")
run(p, contact, size=CONT_S)

# ── PROFESSIONAL SUMMARY ─────────────────────────────────────────
section_head(doc, "PROFESSIONAL SUMMARY", sb=8)
p = para(doc, sa=6)
summary = (
    "M.S. graduate in Information Technology Management (Data Analytics & AI) from "
    "UW–Milwaukee, currently heading all Agentic AI development for the Excelsis360 "
    "education platform. Skilled in machine learning, autonomous agent design, and "
    "full-stack AI application development using Python, LangChain, React, and Snowflake."
)
run(p, summary)

# ── WORK EXPERIENCE ──────────────────────────────────────────────
section_head(doc, "WORK EXPERIENCE")

exp_content = [
    {
        "company": "Excellerate Education Solutions",
        "location": "IL",
        "role": "Software Engineer Intern (Agentic AI)",
        "start": "2026-01", "end": None, "current": True,
        "bullets": [
            "Head all Agentic AI initiatives for Excelsis360, architecting autonomous agent systems that automate complex, multi-step educational workflows end-to-end",
            "Architected and built the Excelsis Data Agent — a full-stack AI analyst with a LangGraph ReAct agent (13 tools), ChromaDB RAG layer (Ollama embeddings) for grounded schema and policy answers, SSE-streaming FastAPI backend, and a React + Tailwind frontend deployed on the Excelsis360 platform",
            "Designed a configurable, domain-agnostic data pipeline connecting SQL Server to a local LLM via natural-language queries, a rate-limited REST API, and an MCP server for Claude Code integration",
        ],
    },
    {
        "company": "Carroll University",
        "location": "Waukesha, WI",
        "role": "Help Desk Technician",
        "start": "2023-09", "end": "2024-05", "current": False,
        "bullets": [
            "Provided first-level technical support and end-user training for campus software applications",
            "Managed help desk ticket queue; configured workstations and assisted in system upgrade deployments",
        ],
    },
    {
        "company": "Cardinal Stritch University",
        "location": "Milwaukee, WI",
        "role": "Help Desk Technician",
        "start": "2021-05", "end": "2023-05", "current": False,
        "bullets": [
            "Trained and supervised student workers; served as escalation point for complex technology issues",
        ],
    },
]

for i, exp in enumerate(exp_content):
    end      = "Present" if exp["current"] else fmt_date(exp["end"])
    date_str = f"{fmt_date(exp['start'])} – {end}"
    label    = f"{exp['role']}  |  {exp['company']}, {exp['location']}"
    sb = 7 if i > 0 else 0
    right_tab_para(doc, label, date_str, lb=True, sb=sb, sa=0)
    for b in exp["bullets"]:
        bullet(doc, b)

# ── EDUCATION ────────────────────────────────────────────────────
section_head(doc, "EDUCATION")

edu_entries = [
    ("M.S., Information Technology Management — Data Analytics & AI",
     "University of Wisconsin – Milwaukee", "Dec. 2025"),
    ("B.S., Business Administration — Communication Minor",
     "Carroll University", "May 2024"),
]

for i, (degree, institution, date) in enumerate(edu_entries):
    sb = 6 if i > 0 else 0
    right_tab_para(doc, degree, date, lb=True, sb=sb)
    p = para(doc)
    run(p, institution, italic=True)

# ── SKILLS ───────────────────────────────────────────────────────
section_head(doc, "SKILLS")

skill_lines = [
    ("Programming Languages", "Python, TypeScript, JavaScript, SQL, HTML, CSS"),
    ("Agentic AI",            "LangChain, LangGraph, ChromaDB, MCP, Ollama"),
    ("Frontend & APIs",       "React, Tailwind CSS, FastAPI, Streamlit"),
    ("Machine Learning",      "scikit-learn, Random Forest, SMOTE, GridSearchCV, Pandas, NumPy"),
    ("Cloud & BI",            "Snowflake, SQL Server, Power BI, Excel, Microsoft Access"),
]
for category, items in skill_lines:
    p = para(doc)
    run(p, category + ": ", bold=True)
    run(p, items)

# ── PROJECTS ─────────────────────────────────────────────────────
section_head(doc, "PROJECTS")

projects = [
    {
        "name":   "Churn Predictor Bot",
        "tech":   "Python, scikit-learn, Streamlit, SMOTE, Pandas",
        "github": "github.com/SusanoKevin/Churn-Predictor",
        "bullet": "Random Forest classifier with GridSearchCV and SMOTE achieving 80%+ accuracy; deployed as an interactive Streamlit app for real-time churn prediction and visualization.",
    },
    {
        "name":   "Excelsis Data Agent",
        "tech":   "Python, LangChain, LangGraph, FastAPI, React, TypeScript, SQL Server, ChromaDB, Ollama, MCP",
        "github": "github.com/SusanoKevin/Excelsis-Data-Analyst-Agent",
        "bullet": "Full-stack AI data analyst for Excelsis360: LangGraph ReAct agent (13 tools), ChromaDB RAG layer for grounded answers, SSE-streaming FastAPI backend, React + Tailwind frontend, configurable SQL Server data backend.",
    },
]

for i, proj in enumerate(projects):
    sb = 6 if i > 0 else 0
    p = para(doc, sb=sb)
    run(p, proj["name"] + "  ", bold=True)
    run(p, f"({proj['tech']})", italic=True, size=9.5)
    p2 = para(doc)
    p2.paragraph_format.left_indent = Inches(0.05)
    run(p2, proj["github"], italic=True, size=9.5)
    bullet(doc, proj["bullet"])

# ── CERTIFICATIONS ───────────────────────────────────────────────
section_head(doc, "CERTIFICATIONS")
right_tab_para(doc,
               "Snowflake Platform Training  —  Snowflake, Inc.",
               "Jun. 2025", lb=False)

# ── Save ─────────────────────────────────────────────────────────
out = ("/Users/susanokevinamalamahilmaran/Downloads/Projects/Resume/"
       "Resume_Susano Kevin Amala_032026 (1).docx")
doc.save(out)
print(f"Saved: {out}")
